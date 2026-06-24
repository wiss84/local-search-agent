"""
DOCX parser for the Local Search Agent ingestion pipeline.

Uses Docling (IBM) for high-quality DOCX extraction:
- Preserves heading hierarchy → Markdown #/##/###
- Converts tables to Markdown
- Handles nested lists, bold/italic
- Strips embedded images (text only)

Large-file protection: DOCX files whose estimated Markdown output exceeds
DOCX_CHAR_SPLIT_THRESHOLD characters are split into paragraph-based batches.
Each batch is converted to a temporary sub-DOCX and fed to docling individually,
capping peak memory to one batch at a time.

Install: pip install "docling>=2.0.0" "python-docx>=1.1.0"
"""

from __future__ import annotations

import io
import logging
import os
import tempfile
from typing import Optional

from local_search_agent.core.document_node import DocumentNode
from local_search_agent.core.key_manager import get_effective_constants
from local_search_agent.ingestion.cleaner import clean
from local_search_agent.ingestion.parser import BaseParser, ParserError

logger = logging.getLogger(__name__)

# How many paragraphs to sample when estimating total output size.
_DOCX_SAMPLE_PARAGRAPHS = 10


# --------------------------------------------------------------------------- #
# Internal helpers                                                              #
# --------------------------------------------------------------------------- #


def _estimate_docx_output_chars(filename: str) -> Optional[int]:
    """
    Estimate the total Markdown character count that docling will produce
    for a DOCX file.

    Samples up to *_DOCX_SAMPLE_PARAGRAPHS* paragraphs with python-docx,
    measures their raw text length, and extrapolates across the total
    paragraph count.

    Returns ``None`` if python-docx is unavailable or statistics fail.
    """
    try:
        from docx import Document  # noqa: PLC0415
    except ImportError:
        logger.debug("python-docx not available; cannot estimate docx size.")
        return None

    try:
        doc = Document(filename)
    except Exception as e:
        logger.debug("python-docx failed to open %r: %s", filename, e)
        return None

    all_paras = list(doc.paragraphs)
    all_tables = doc.tables

    total_paras = len(all_paras)
    total_tables = len(all_tables)

    if total_paras == 0 and total_tables == 0:
        return 0

    # Average raw text length of sampled paragraphs.
    sample = all_paras[:_DOCX_SAMPLE_PARAGRAPHS]
    sample_text = "".join(p.text for p in sample)
    sample_count = len(sample)
    avg_para_chars = len(sample_text) / sample_count if sample_count > 0 else 0.0

    # Rough heuristic: each table row (as text) averages ~80 chars; each table
    # averages 20 rows → ~1 600 chars per table including header/labels.
    avg_table_chars = 1_600.0

    # Projected total output chars.
    para_total = total_paras * avg_para_chars
    table_total = total_tables * avg_table_chars
    estimated = int(para_total + table_total)

    logger.debug(
        "DOCX size estimate for %r: ~%d chars  (paras=%d × %.0f avg + tables=%d × ~%.0f)",
        filename,
        estimated,
        total_paras,
        avg_para_chars,
        total_tables,
        avg_table_chars,
    )
    return estimated


def _build_sub_docx(
    paragraphs: list,
    tables: list,
) -> Optional[bytes]:
    """
    Build a valid .docx binary (in-memory) containing only *paragraphs* and
    *tables*, using python-docx as the builder.

    Returns ``None`` on failure.
    """
    try:
        import copy  # noqa: PLC0415

        from docx import Document  # noqa: PLC0415
    except ImportError:
        return None

    buf = io.BytesIO()
    new_doc = Document()

    for para in paragraphs:
        try:
            new_doc._body._element.append(copy.deepcopy(para._element))  # type: ignore[attr-defined]
        except Exception:
            # If deep-copy fails (e.g. HTML altChunk), write text only.
            new_doc.add_paragraph(para.text)

    for table in tables:
        try:
            new_doc._body._element.append(copy.deepcopy(table._element))  # type: ignore[attr-defined]
        except Exception:
            pass

    try:
        new_doc.save(buf)
        return buf.getvalue()
    except Exception as e:
        logger.debug("Failed to build sub-DOCX: %s", e)
        return None


def _split_docx_in_batches(
    docx_bytes: bytes,
    *,
    source_path: str,
    char_split_threshold: int,
) -> str:
    """
    Split a large DOCX into paragraph-count batches and convert each with
    docling independently.

    Parameters
    ----------
    docx_bytes : The full original DOCX binary.
    source_path: Original source filename (for error messages).

    Returns Markdown string (concatenated across batches).
    """
    from docling.document_converter import DocumentConverter  # noqa: PLC0415
    from docx import Document  # noqa: PLC0415

    try:
        src = Document(io.BytesIO(docx_bytes))
    except Exception as e:
        raise ParserError(source_path, f"python-docx failed to parse input: {e!r}", original=e)

    all_paragraphs: list = list(src.paragraphs)
    all_tables: list = list(src.tables)
    total_paras = len(all_paragraphs)

    if total_paras == 0:
        # No paragraphs — let docling handle whatever structure exists.
        raise ParserError(source_path, "python-docx found no paragraphs (invalid or empty?).")

    # --- Determine batch size in paragraphs ---
    # We estimate the output char count from a sample, divide the total
    # expected chars by the threshold, then ceiling to batches count.
    est_total_chars = _estimate_docx_output_chars(source_path) or 0

    if est_total_chars <= char_split_threshold:
        # Under threshold — single call
        converter = DocumentConverter()
        result = converter.convert(source_path)
        return result.document.export_to_markdown()

    num_batches = max(2, (est_total_chars // char_split_threshold) + 1)
    paras_per_batch = max(1, total_paras // num_batches)
    logger.info(
        "Large DOCX (%d paras, ~%d est. chars) — splitting to %d batches (~%d paras each)",
        total_paras,
        est_total_chars,
        num_batches,
        paras_per_batch,
    )

    converter = DocumentConverter()
    accumulated: list[str] = []
    warn_parts: list[str] = []

    batch_num = 0
    for start in range(0, total_paras, paras_per_batch):
        end_para = min(start + paras_per_batch, total_paras)
        batch_paras = all_paragraphs[start:end_para]

        # Tables are only included in the first batch to preserve their
        # position context.  Tables in later batches would be orphaned
        # without surrounding paragraphs which are difficult to match
        # reliably using python-docx XML indices.
        batch_tables = all_tables if start == 0 else []

        sub_bytes = _build_sub_docx(batch_paras, batch_tables)
        if sub_bytes is None:
            msg = f"[batch {batch_num + 1}] skipped: could not build sub-DOCX"
            warn_parts.append(msg)
            logger.warning("%s", msg)
            batch_num += 1
            continue

        tf = tempfile.NamedTemporaryFile(suffix=".docx", delete=False, prefix="docx_batch_")
        tf.write(sub_bytes)
        tf.close()

        try:
            result = converter.convert(tf.name)
            batch_md = result.document.export_to_markdown()
            accumulated.append(batch_md)
            logger.debug(
                "Converted DOCX batch %d  (paras %d-%d)",
                batch_num + 1,
                start + 1,
                end_para,
            )
        except MemoryError as e:
            msg = (
                f"[batch {batch_num + 1}] skipped: out of memory ({e!r}). "
                "Consider reducing DOCX_CHAR_SPLIT_THRESHOLD."
            )
            warn_parts.append(msg)
            logger.warning("%s", msg)
        except Exception as e:
            msg = f"[batch {batch_num + 1}] skipped: conversion error: {e!r}"
            warn_parts.append(msg)
            logger.warning("%s", msg)
        finally:
            os.unlink(tf.name)

        batch_num += 1

    if not accumulated:
        raise ParserError(
            source_path,
            "All DOCX batches failed. The document may be empty or corrupted.",
        )

    combined = "\n\n".join(accumulated)
    if warn_parts:
        combined += "\n\n_" + "  ".join(warn_parts) + "_"

    return combined


# --------------------------------------------------------------------------- #
# Parser                                                                       #
# --------------------------------------------------------------------------- #


class DOCXParser(BaseParser):
    """
    Parse DOCX files using Docling.

    Docling handles:
    - Heading levels → Markdown headings
    - Tables → Markdown tables
    - Numbered and bulleted lists → Markdown lists
    - Bold/italic inline formatting preserved

    Large files are transparently split into paragraph batches so that
    docling's peak memory footprint stays bounded.
    """

    @property
    def supported_extensions(self) -> frozenset[str]:
        return frozenset({".docx", ".doc"})

    def parse(
        self,
        source_path: str,
        workspace: str,
        title: Optional[str] = None,
    ) -> DocumentNode:
        if not os.path.isfile(source_path):
            raise FileNotFoundError(f"DOCX file not found: {source_path!r}")

        try:
            from docling.document_converter import DocumentConverter
        except ImportError as e:
            raise ParserError(
                source_path,
                "Docling is not installed. Run: pip install 'docling>=2.0.0'",
                original=e,
            )

        try:
            import docx  # noqa: F401 — ensures python-docx is importable
        except ImportError as e:
            raise ParserError(
                source_path,
                "python-docx is not installed. Run: pip install 'python-docx>=1.1.0'",
                original=e,
            )

        logger.info("Parsing DOCX: %s", source_path)

        estimated_chars: Optional[int] = None
        use_batching = False

        # Estimate output size safely — failures fall through to single-call.
        try:
            import docx  # noqa: F401 — checks python-docx availability

            estimated_chars = _estimate_docx_output_chars(source_path)
            docx_char_split_threshold = get_effective_constants()["DOCX_CHAR_SPLIT_THRESHOLD"]
            use_batching = bool(
                estimated_chars is not None and estimated_chars > docx_char_split_threshold
            )
        except ImportError as e:
            raise ParserError(
                source_path,
                "python-docx is not installed. Run: pip install 'python-docx>=1.1.0'",
                original=e,
            )

        try:
            from docling.document_converter import DocumentConverter

            converter = DocumentConverter()

            if use_batching:
                # Read the original bytes so we can operate on a copy.
                with open(source_path, "rb") as fh:
                    file_bytes = fh.read()
                raw_markdown = _split_docx_in_batches(
                    file_bytes,
                    source_path=source_path,
                    char_split_threshold=docx_char_split_threshold,
                )
            else:
                result = converter.convert(source_path)
                raw_markdown = result.document.export_to_markdown()

        except ParserError:
            raise
        except Exception as e:
            raise ParserError(source_path, f"Docling conversion failed: {e}", original=e)

        cleaned_text = clean(raw_markdown)

        return DocumentNode.from_file(
            source_path=source_path,
            text=cleaned_text,
            workspace=workspace,
            title=title,
        )
