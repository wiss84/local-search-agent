"""
Word (.docx) chat export.

Converts the structured message list from a chat session (the same shape
returned by GET /sessions/{id}/messages) into a formatted .docx document:
"You" / "Assistant" headings, paragraphs, fenced code blocks in monospace,
and citation links.

Citation links in agent answers look like:
    [Document Title](http://localhost:8000/docs/<doc_id>)

That URL only resolves while the dashboard's file server is running. For an
exported document that should still be useful after the app is closed, we
resolve doc_id -> the document's actual source_path via workspace_manager
and emit a real file:// hyperlink instead. If resolution fails (file moved
or deleted since the chat happened), we fall back to the original URL
rather than dropping the link.

This is intentionally a light-touch renderer, not a full Markdown parser:
it only needs to handle what the agent's own output actually produces
(plain paragraphs, a trailing "Sources:" list, occasional fenced code
blocks) — not arbitrary Markdown.
"""

from __future__ import annotations

import io
import logging
import pathlib
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

logger = logging.getLogger(__name__)

_CODE_BLOCK_RE = re.compile(r"```(?:\w+)?\n(.*?)```", re.DOTALL)
_CITATION_LINK_RE = re.compile(r"\[([^\]]+)\]\((\S+/docs/([a-f0-9]+))\)")


def _resolve_citation_uri(doc_id: str, workspace_manager) -> str | None:
    """Return a stable file:// URI for doc_id, or None if it can't be resolved."""
    try:
        node = workspace_manager.get_document(doc_id)
    except Exception:
        return None
    if node is None or not getattr(node, "source_path", None):
        return None
    path = pathlib.Path(node.source_path)
    if not path.is_file():
        return None
    return path.as_uri()


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    """
    Insert a real clickable hyperlink run into a python-docx paragraph.

    python-docx has no public API for hyperlinks, so this manipulates the
    underlying XML relationship directly. This is a well-known, stable
    pattern for python-docx (not fragile, just undocumented in the public
    API) — it does not touch any private/unstable internals beyond the
    documented OOXML relationship mechanism.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2563EB")
    rpr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)

    new_run.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _render_text_block(doc: Document, text: str, workspace_manager) -> None:
    """Render plain (non-code-block) text as paragraphs, with citation links as real hyperlinks."""
    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if not line:
            continue

        p = doc.add_paragraph()
        last_end = 0
        for m in _CITATION_LINK_RE.finditer(line):
            if m.start() > last_end:
                p.add_run(line[last_end : m.start()])
            label, raw_url, doc_id = m.group(1), m.group(2), m.group(3)
            resolved = _resolve_citation_uri(doc_id, workspace_manager)
            _add_hyperlink(p, resolved or raw_url, label)
            last_end = m.end()
        if last_end < len(line):
            p.add_run(line[last_end:])


def _render_message_body(doc: Document, content: str, workspace_manager) -> None:
    """Render one message's content: text blocks interleaved with fenced code blocks."""
    pos = 0
    for m in _CODE_BLOCK_RE.finditer(content):
        _render_text_block(doc, content[pos : m.start()], workspace_manager)

        code_text = m.group(1).rstrip("\n")
        p = doc.add_paragraph()
        run = p.add_run(code_text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        pos = m.end()

    _render_text_block(doc, content[pos:], workspace_manager)


def build_docx(messages: list[dict], workspace_manager) -> bytes:
    """
    Build a .docx document from a list of {"role": "user"|"assistant", "content": str}
    dicts and return the raw file bytes.
    """
    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)

    for msg in messages:
        doc.add_heading("You" if msg.get("role") == "user" else "Assistant", level=2)
        _render_message_body(doc, msg.get("content", "") or "", workspace_manager)
        doc.add_paragraph()  # spacing between turns

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
